from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path('docs')
OUT.mkdir(exist_ok=True)
INK = '0B1121'
BLUE = '1F4D78'
MUTED = '5B6472'
LINE = 'D9DEE7'
FILL = 'EEF2F7'
PALE = 'F7F9FC'
WHITE = 'FFFFFF'


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for side, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{side}'))
        if node is None:
            node = OxmlElement(f'w:{side}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def set_table_widths(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in('w:tblW')
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), str(sum(widths)))
    tbl_w.set(qn('w:type'), 'dxa')
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn('w:w'), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in('w:tcW')
            if tc_w is None:
                tc_w = OxmlElement('w:tcW')
                tc_pr.append(tc_w)
            tc_w.set(qn('w:w'), str(width))
            tc_w.set(qn('w:type'), 'dxa')
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def put_text(cell, value, bold=False, color=INK, size=9.5):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(str(value))
    r.bold = bold
    r.font.name = 'Calibri'
    r._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    r._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)


def table(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.style = 'Table Grid'
    for cell, text in zip(t.rows[0].cells, headers):
        set_cell_shading(cell, FILL)
        put_text(cell, text, bold=True, color=BLUE, size=9)
    for index, row in enumerate(rows):
        cells = t.add_row().cells
        for cell, value in zip(cells, row):
            if index % 2:
                set_cell_shading(cell, PALE)
            put_text(cell, value)
    set_table_widths(t, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def configure(doc, label):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.right_margin = Inches(1.0)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(1.0)
    sec.header_distance = Inches(0.45)
    sec.footer_distance = Inches(0.45)
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Calibri'
    normal._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, colour, before, after in [
        ('Heading 1', 16, BLUE, 16, 8),
        ('Heading 2', 13, BLUE, 12, 6),
        ('Heading 3', 11.5, INK, 9, 4),
    ]:
        style = styles[name]
        style.font.name = 'Calibri'
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(colour)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for list_name in ('List Bullet', 'List Number'):
        style = styles[list_name]
        style.font.name = 'Calibri'
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.1
    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run(f'KRONIQOS  |  {label}')
    run.font.name = 'Calibri'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run('Confidential — Product planning')
    run.font.name = 'Calibri'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)


def heading(doc, text, level=1):
    doc.add_paragraph(text, style=f'Heading {level}')


def para(doc, text, lead=None):
    p = doc.add_paragraph()
    if lead and text.startswith(lead):
        r = p.add_run(lead)
        r.bold = True
        p.add_run(text[len(lead):])
    else:
        p.add_run(text)
    return p


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')


def numbers(doc, items):
    for item in items:
        doc.add_paragraph(item, style='List Number')


def cover(doc, kind, subtitle):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(76)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(kind.upper())
    r.bold = True; r.font.name = 'Calibri'; r.font.size = Pt(10); r.font.color.rgb = RGBColor.from_string(BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run('Kroniqos')
    r.bold = True; r.font.name = 'Calibri'; r.font.size = Pt(30); r.font.color.rgb = RGBColor.from_string(INK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run(subtitle)
    r.font.name = 'Calibri'; r.font.size = Pt(15); r.font.color.rgb = RGBColor.from_string(MUTED)
    table(doc, ['Document control', 'Value'], [
        ('Product', 'Kroniqos — multi-tenant business operations platform'),
        ('Prepared for', 'Product Management'),
        ('Version', '1.0'),
        ('Date', date.today().strftime('%d %B %Y')),
        ('Status', 'Current implemented scope and operating specification'),
        ('Technology', 'React, TypeScript, Ant Design, Tailwind, Zustand, Dexie, Supabase, Edge Functions'),
    ], [2700, 6660])
    para(doc, 'This document reflects the implemented product scope. It is intended to support planning, demos, acceptance testing, training, and future release decisions.')
    doc.add_page_break()


def make_prd():
    doc = Document(); configure(doc, 'Product Requirements Document')
    cover(doc, 'Product requirements document', 'A SaaS platform for retail, service businesses, and online selling.')
    heading(doc, '1. Executive summary')
    para(doc, 'Kroniqos is a multi-tenant SaaS application for African businesses that sell products, deliver services, or do both. It combines touch-friendly point of sale, inventory and cost tracking, customer credit, expenses, projects/contracts, an optional public storefront, and a platform administration layer.')
    para(doc, 'The product is built around a pragmatic operating model: the checkout remains offline-capable to protect customer service during unreliable connectivity, while shared financial, inventory, management, and platform controls are online-first so the server remains the source of truth.')
    heading(doc, '2. Product vision, objectives, and boundaries')
    table(doc, ['Area', 'Requirement / success outcome'], [
        ('Vision', 'Enable a small or growing business to run day-to-day operations and understand its cash, stock, customers, and service work from one system.'),
        ('Primary market', 'African retail and service businesses, initially Nigerian operations using ₦ but with organisation-level currency selection.'),
        ('Business model', 'Multi-tenant SaaS with Starter, Growth, and higher plans; plans control modules, product limits, storefront, branding, and capabilities.'),
        ('MVP operating shape', 'Single-store tenant is fully supported. Store IDs and organisation boundaries allow branches to be added later.'),
        ('Primary outcomes', 'Fast checkout, auditable stock, visible margin, controlled staff access, payment/credit visibility, and self-service onboarding.'),
    ], [2250, 7110])
    heading(doc, 'In scope', 2)
    bullets(doc, ['Tenant onboarding, user authentication, organisation setup, plans and entitlements.', 'Retail POS, stock and supplier deliveries, cash shifts, sales/returns, expenses, customer credit, reports and CSV import/export.', 'Service catalogue, clients, projects/contracts, project payments, stages, expenses, staff assignments, documents and invoices.', 'Growth-plan public storefront with product catalogue, online ordering foundation, content sections, brand colour/logo/hero images, and about/contact pages.', 'Platform operations for organisations, plans, billing configuration, audits, analytics, support/safety, maintenance status, and platform team RBAC.'])
    heading(doc, 'Out of current scope', 2)
    bullets(doc, ['Multi-branch stock transfers, warehouse replenishment automation, purchase orders, accounting integration, loyalty, and automated tax filing.', 'A built-in payment terminal driver or cash-drawer hardware integration; payment methods are recorded operationally.', 'Full marketplace ordering, delivery logistics, refunds through a payment gateway, or customer accounts for the public storefront.', 'Self-service enterprise SSO, multi-currency settlement, and formal accounting ledger.'])
    heading(doc, '3. Personas and access model')
    table(doc, ['Persona', 'Primary needs', 'Access boundary'], [
        ('Tenant administrator', 'Set up company, staff, plans/settings, products, pricing, reports, credit and governance.', 'Full tenant administration; cannot administer other organisations or platform controls.'),
        ('Manager', 'Maintain daily operations: inventory, deliveries, returns, expenses, sales review, service work.', 'Operational authority, with restricted staff/platform configuration and profit access where configured.'),
        ('Cashier', 'Serve customers quickly, take payments, hold/resume sales, operate own shift.', 'Checkout-focused; cannot change shared stock, pricing, staff, or management reporting.'),
        ('Service operator', 'Create and deliver projects, update stages, log payments/expenses, send client update.', 'Service module access according to tenant-assigned role.'),
        ('Organisation owner', 'Own business account, plan, subscription and high-level outcomes.', 'May also be tenant admin; organisation-wide business responsibilities.'),
        ('Platform owner/operator/support/finance/viewer', 'Manage SaaS organisations, plans, billing, support and platform insights.', 'Platform RBAC is separate from tenant roles and is server-enforced.'),
        ('Public shopper', 'Discover a tenant’s online catalogue and place an order.', 'Public storefront only; no tenant dashboard access.'),
    ], [1550, 3450, 4360])
    heading(doc, '4. Product modules and requirements')
    table(doc, ['Module', 'Requirements'], [
        ('Authentication & onboarding', 'Email/password sign-in; modern tenant/platform login; company registration; first administrator account; guided organisation setup; optional starter-data setup tools.'),
        ('Tenant dashboard', 'Unified overview with module tabs for retail and services; date period filters; business stats, financial movement, payment mix and expense visibility.'),
        ('Checkout', 'Barcode/SKU/name search; touch-first quantity modal; cash/card/transfer/credit; default cash tender equal to total; change calculation; held sales; printable receipt; customer display on second screen.'),
        ('Historical sales', 'Manager/admin can record past sale date, original price, payment method and optional stock deduction. Historical records appear distinctly in sales/reporting.'),
        ('Inventory & catalogue', 'Products, categories, barcode/SKU, cost/sell price, low-stock threshold, count, adjustments, movement log, product limits and price history.'),
        ('Deliveries & warehouse', 'Supplier delivery date, supplier, location, quantity and cost. Optional warehouse/shop-floor locations and stock location tracking.'),
        ('Sales, returns & shifts', 'Sales filters, CSV export/import, receipt item details, partial item return/exchange online, cash shift opening/closing/variance and return audit.'),
        ('Credit', 'Credit checkout requires customer name/phone, optional due date and initial payment; register tracks balance, status, partial repayments, payment dates and printable history.'),
        ('Expenses', 'Categorised expense register, custom categories, date, CSV bulk import with warning guidance and sample download.'),
        ('Services', 'Independent service types with description/benchmark price; client capture; project wizard; contract amount/deposit/dates/documents; workflow stages, notes, staff, payments, expenses and invoice print.'),
        ('Storefront', 'Growth-plan-controlled online storefront with store/company branding, hero, sections, featured/recent/categories, product image limit, descriptions, detail pages, search and static content.'),
        ('Platform control', 'Responsive platform dashboard, organisation operations, plan editor/entitlements, billing operations, analytics, platform audit, support/safety, status and platform team.'),
    ], [2100, 7260])
    heading(doc, '5. Key product decisions')
    table(doc, ['Decision', 'Rationale', 'Implication'], [
        ('Checkout offline-first', 'Tills must keep serving customers during weak network.', 'Cart, held sales and queued checkout use browser storage; background sync confirms server state later.'),
        ('All other shared operations online-first', 'Stock adjustments, deliveries, prices, returns, staff and settings need one authoritative truth.', 'Actions fail safely when offline rather than creating conflicting local state.'),
        ('Products support cost snapshots', 'Cost prices change over time.', 'Profit reporting uses sale-time cost, not current catalogue cost.'),
        ('Credit is a sale method', 'Businesses often sell now and collect in parts later.', 'Creates a customer credit ledger with initial and subsequent payments.'),
        ('Retail and services can coexist', 'Many businesses sell products and perform paid services.', 'Organisation business modes enable tailored menus and dashboard tabs rather than separate products.'),
        ('Storefront is entitlement-based', 'Public commerce adds storage, content and support requirements.', 'Online storefront navigation is hidden unless Growth-level entitlement allows it.'),
    ], [1850, 3750, 3760])
    heading(doc, '6. Success measures')
    table(doc, ['Objective', 'Measure', 'Initial target'], [
        ('Checkout resilience', 'Offline transactions retained and synced exactly once after reconnect.', '100% in controlled outage test.'),
        ('Checkout efficiency', 'Normal cashier flow from scan to completed cash sale.', 'Under 30 seconds for a small basket after training.'),
        ('Inventory integrity', 'Delivery/count/return actions generate a server movement/audit record.', '100% of authorised operations.'),
        ('Adoption', 'New organisation completes onboarding and reaches usable dashboard.', 'Within 15 minutes without PM intervention.'),
        ('Financial visibility', 'Admin can view filtered revenue, COGS, gross profit and expenses.', 'Available by day, month and year.'),
        ('SaaS management', 'Platform team can inspect tenant status/plan/metrics/audit without database access.', 'All active organisations.'),
    ], [2400, 4300, 2660])
    heading(doc, '7. Non-functional requirements')
    bullets(doc, ['Mobile-first responsive layout with tenant bottom navigation; responsive platform off-canvas mobile sidebar; touch targets are intentionally large at checkout.', 'Light theme default and dark mode supported. UI uses a black/white/charcoal design system with #0B1121 as primary; status colours communicate warning, success, and error.', 'Supabase Auth, Row Level Security and server-side RPC/Edge Function checks protect tenant data. Client UI visibility is not the security boundary.', 'All monetary values are displayed using the selected organisation currency; initial Nigerian tenant flow uses ₦ and database amount calculations use integer kobo where applicable.', 'Loading states, actionable errors, warning notices and sync status must explain what happened without hiding business data.', 'The app must preserve queued checkout data through page reload; reset tools must protect unsynced business records.'])
    heading(doc, '8. Dependencies, risks, and release criteria')
    table(doc, ['Risk / dependency', 'PM action'], [
        ('Supabase migrations and Edge Functions', 'Keep an ordered release checklist. Validate migrations and deploy functions before tenant acceptance tests.'),
        ('RLS/profile role consistency', 'Create role test accounts and validate both UI gating and server responses for every critical action.'),
        ('Paystack configuration', 'Billing cannot move subscriptions automatically until required keys, webhook and callback configuration are completed.'),
        ('Offline sync conflict', 'Train stores on reconciliation. Do not allow destructive local resets with pending checkout queue.'),
        ('Plan enforcement', 'Define product/storage/staff/storefront/service limits before public launch and test plan downgrade behaviour.'),
        ('Customer data', 'Publish privacy, retention, backup, support and incident policies before broad SaaS onboarding.'),
    ], [3200, 6160])
    heading(doc, '9. Release acceptance checklist')
    bullets(doc, ['Tenant administrator can register a company, create/activate initial configuration and sign in.', 'Cashier completes, holds, resumes and syncs a sale; credit sale clears cart/profile state and defaults back to cash on the next transaction.', 'Manager/admin can receive delivery, count/adjust stock, add products/categories and see audited online changes.', 'Admin can review sales, quantity/item receipts, returns, credit payments, expenses and gross-profit reporting.', 'Service-enabled tenant can create a client/project, upload document, record payment/expense/stage and print tenant-branded invoice.', 'Growth tenant can publish online storefront content; Starter tenant cannot see the storefront module.', 'Platform role can view organisation records, edit entitlement plans, review audit data and open responsive dashboard on mobile.', 'All date controls use Ant Design date picker/range picker and all CSV imports use warning guidance plus sample download.'])
    doc.save(OUT / 'Kroniqos_PRD.docx')


def make_functional_spec():
    doc = Document(); configure(doc, 'Functional Specification')
    cover(doc, 'Functional specification', 'System behaviour, user flows, roles, data responsibilities and acceptance rules.')
    heading(doc, '1. System context')
    para(doc, 'Kroniqos contains two protected application surfaces and one public surface: the tenant operations application, the Kroniqos platform control application, and an optional tenant storefront. All tenant data is scoped to an organisation and normally to one or more stores.')
    table(doc, ['Surface', 'Audience', 'Purpose'], [
        ('Tenant application', 'Tenant staff', 'Run retail, service, finance and configuration operations for one organisation.'),
        ('Platform application', 'Kroniqos platform staff', 'Operate organisations, plans, billing, support/safety, status, analytics and platform audit.'),
        ('Public storefront', 'Customers of a Growth tenant', 'Browse tenant catalogue/content and begin ordering without dashboard access.'),
    ], [2100, 2500, 4760])
    heading(doc, '2. Authentication, tenancy and RBAC')
    table(doc, ['Function', 'Behaviour', 'Authorisation'], [
        ('Tenant sign-in', 'User signs in using Supabase Auth email/password and is routed to their tenant workspace when active membership/profile is resolved.', 'Authenticated user with valid active tenant membership.'),
        ('Platform sign-in', 'User signs in at platform route; platform access is checked separately from tenant membership.', 'Platform owner/operator/support/finance/viewer as assigned.'),
        ('Company registration', 'Prospective owner creates account then completes company/organisation onboarding.', 'Unauthenticated registration; email confirmation follows Supabase project settings.'),
        ('Tenant role enforcement', 'Menu visibility reflects role, but RPCs/RLS validate role/store on all protected data mutations.', 'Cashier, manager, admin.'),
        ('Staff lifecycle', 'Admin creates staff using Edge Function. Deactivated/suspended memberships must remove tenant access.', 'Tenant admin for staff actions; service role remains server-side only.'),
        ('Platform role enforcement', 'Platform pages and RPCs enforce platform RBAC; access by manually typing a URL must not provide data.', 'Platform role required per operation.'),
    ], [1750, 5000, 2610])
    heading(doc, '3. Tenant navigation and responsive behaviour')
    bullets(doc, ['Desktop tenant navigation is collapsible; content width reflows when it opens or closes.', 'Mobile tenant navigation uses a bottom bar for Home, Checkout, Cart, relevant module shortcut, and More. More opens the sidebar; tapping outside closes it.', 'Mobile header uses a three-dot account menu for user name, role and sign out.', 'Platform desktop sidebar stays visible. On mobile it becomes an off-canvas drawer opened by the header hamburger; selection or outside tap dismisses it.', 'The footer occupies the bottom of short pages without overlaying content.'])
    heading(doc, '4. Retail checkout functional flow')
    numbers(doc, ['Cashier opens Checkout. The local product catalogue is shown in a paginated/searchable product list; cart begins empty and payment defaults to Cash.', 'Cashier searches by name, SKU or barcode. A recognised scanner input with Enter adds the product. Selecting a product opens a quantity entry modal so a large quantity can be entered directly.', 'Cart displays line quantity, unit price and totals. Cashier may change quantity; authorised users may apply discount or, in historical mode, set original sale price.', 'Cashier selects Cash, Card/POS, Transfer or Credit. Cash tendered begins at the exact total; user changes it only when change is required.', 'Credit opens a modal requiring customer name and phone and optionally accepts initial payment and expected payment date. Credit amount is sale total minus initial payment.', 'Cashier completes sale. Online checkout is confirmed against server; offline-compatible checkout is stored locally with durable client ID and added to sync outbox.', 'Cart, customer credit details, cash tender and payment method reset after a successful completion. The next sale defaults to Cash.', 'Receipt is viewable/printable. On a dual-screen machine the customer-display action opens the customer display view in a separate window; it must not replace the cashier tab.'])
    table(doc, ['Checkout exception', 'Required behaviour'], [
        ('No matching barcode', 'Show not-found message. Manager/admin can create product with scanned barcode prefilled; cashier cannot create catalogue item.'),
        ('Insufficient cash tender', 'Disable completion and clearly show amount remaining.'),
        ('Offline', 'Keep selling from cached catalogue. Sync badge shows pending count. Do not silently drop records.'),
        ('Sync rejected', 'Keep queue record and show actionable reason. Typical examples include role, product ID, price/total validation or stock validation.'),
        ('Hold sale', 'Persist locally, clear active cart, make sale visible directly below completion controls, resume restores exact cart; discard is explicit.'),
        ('Historical sale', 'Manager/admin-only date picker. Past sale does not deduct current stock unless stock correction is explicitly checked; it is labelled in history.'),
    ], [3050, 6310])
    heading(doc, '5. Inventory, catalogue and deliveries')
    table(doc, ['Use case', 'Inputs', 'System result'], [
        ('Add product', 'Barcode/SKU first, name, description, category, cost price, selling price, stock/low-stock threshold.', 'Online server creates store-scoped product; product limit is checked; local catalogue refreshes.'),
        ('Edit product', 'Product fields, online availability, up to two product images where storefront entitlement permits.', 'Online update persists to Supabase and product/storefront reads refresh.'),
        ('Category', 'Existing category or free-text category when adding product.', 'Category is created/reused online and assigned to product.'),
        ('Receive supplier delivery', 'Product, stock location, supplier, quantity, unit cost, received date.', 'Online inventory movement increases location stock; delivery log is created; cost price and price-history audit are updated as applicable.'),
        ('Count stock', 'Physical count.', 'Online adjustment records variance and count movement.'),
        ('Adjust stock', 'Quantity delta / desired count and reason.', 'Online inventory ledger records adjustment and actor; local product cache refreshes.'),
        ('Warehouse', 'Optional named locations including shop floor and warehouse.', 'Stock can be received into a location; checkout stock is expected from shop floor.'),
    ], [2100, 3500, 3760])
    heading(doc, '6. Sales, returns, shifts, reports and imports')
    table(doc, ['Feature', 'Functional rule'], [
        ('Sales list', 'Supports date and payment/cashier filters, receipt details with individual line items and quantities, local/online visibility, CSV export and authorised historical bulk upload.'),
        ('Returns/exchanges', 'Online manager/admin workflow. User selects original sale, individual line items and quantities; server prevents over-return, restores stock and writes actor/date audit trail.'),
        ('Cash shifts', 'Cashier/manager opens and closes a shift. System calculates expected cash, permits counted cash and variance reason, and keeps an auditable closing record.'),
        ('Profit reports', 'Admin sees period, date/month/year filters and per-line quantity, unit selling price, sale-time unit cost, COGS and profit. All summary values are derived from included sale lines.'),
        ('Bulk CSV', 'Expenses and historical sales offer warning guidance, sample CSV download, file selection, preview and row-level error handling. Import is server-backed and does not bypass roles.'),
        ('Date selection', 'Operational dates use Ant Design DatePicker. Filter periods spanning two dates use RangePicker. Stored server values are formatted YYYY-MM-DD.'),
    ], [2700, 6660])
    heading(doc, '7. Customer credit and expenses')
    table(doc, ['Credit lifecycle', 'Behaviour'], [
        ('Create', 'A credit checkout creates or links customer profile, records total, initial payment if provided, outstanding balance, optional due date and credit status.'),
        ('Register', 'Credit page shows customer, original amount, amount paid, balance, due date and status including paid/active/overdue where calculated.'),
        ('Partial payment', 'Manager/admin records amount and payment date. Server validates entitlement and remaining balance, then appends payment history.'),
        ('Payment evidence', 'Credit detail can show printable payment history for reconciliation and customer communication.'),
        ('Expense creation', 'Manager/admin records category (select or add), description, amount and expense date. Expense is immediately online-authoritative and displayed in reports/dashboard.'),
    ], [2650, 6710])
    heading(doc, '8. Service/project operations')
    numbers(doc, ['Tenant enables Services business mode through platform/organisation configuration. Service-only businesses see service-relevant dashboard/menu; businesses with both modes see both.', 'Admin/manager creates reusable service definitions containing title, description and optional benchmark price.', 'New Project wizard: select existing client or add client/company name and required phone; email/address optional. Then select service, title, description, project date, estimated delivery, value, optional deposit and PDF document.', 'Project details drawer shows client, service, status, stage, dates, quoted amount, received amount and balance.', 'Team can record project payments, choose payment method/date/reference/note, add project expenses, assign staff, attach documents and update workflow stage/comment.', 'Stage activity/history is retained. When configured, stage update can be used to prepare/send a WhatsApp client update.', 'Invoice print includes tenant/company logo and name, contract details, payment breakdown and outstanding balance.'])
    table(doc, ['Project status', 'Interpretation'], [
        ('New / ongoing', 'Project has been created and is moving through stages.'),
        ('Due soon / overdue', 'Derived from estimated delivery date and current date; should be visually clear to operations.'),
        ('Completed', 'Team closes project when delivery/work is done; financial balance may still be visible for collection.'),
    ], [2800, 6560])
    heading(doc, '9. Storefront and content management')
    table(doc, ['Capability', 'Rule'], [
        ('Entitlement', 'Only Growth or higher plan tenants see storefront management/navigation and may activate public shop.'),
        ('Branding', 'Company name is source of storefront name; Growth tenant may upload logo, choose primary colour and use up to two hero images.'),
        ('Products', 'Tenant can add product description, online availability/publish state and up to two product images to protect storage usage.'),
        ('Page layout', 'Public pages include professional header/hero/product listing/footer, product details, About and Contact with mini hero image and breadcrumb.'),
        ('Homepage sections', 'Tenant manages featured/recent product blocks, categories, custom sections, vision/mission, testimonials and calls to action.'),
        ('Search and ordering', 'Storefront provides catalogue search and ordering foundation. It must only show products published/available online for that tenant.'),
    ], [2500, 6860])
    heading(doc, '10. SaaS platform operations')
    table(doc, ['Area', 'Functional behaviour'], [
        ('Organisation operations', 'List/filter organisation, open profile, see stores/status/staff count/owners, change plan, resend supported onboarding/invites, suspend/reactivate where authorised.'),
        ('Plan editor', 'Platform-authorised users update plan name, description, monthly price, active state, numeric limits and feature flags. Changes create platform audit events.'),
        ('Billing', 'Tracks trials, current plan/period, payment history, Paystack subscription configuration and past-due/suspension rules. External billing requires deployed/configured Edge Functions and Paystack secrets.'),
        ('Audit log', 'Search and date range filter system actions such as plan, suspension, settings and operational actions; shows actor, organisation, timestamp and before/after data.'),
        ('Analytics', 'Date range reports active organisations, new sign-ups, churn, stores, staff, sales volume, subscription revenue, trials started and trial-to-paid conversion.'),
        ('Support & safety', 'Support notes/tickets, time-limited/audited organisation view access, maintenance notices and system-status controls.'),
        ('Platform team', 'Platform RBAC dashboard controls owner, operator, support, finance and viewer permissions.'),
    ], [2600, 6760])
    heading(doc, '11. Data, sync and security requirements')
    table(doc, ['Concern', 'Requirement'], [
        ('Tenant boundary', 'Every tenant record is scoped by organisation/store and protected by Supabase RLS; user must not read/write other organisation data.'),
        ('Offline database', 'Dexie/IndexedDB stores checkout-relevant catalogue, cart/held-sale state, sale items, sync outbox and local receipts. It is not the full business system of record.'),
        ('Sync method', 'Background sync pushes durable checkout/outbox items idempotently, then pulls relevant cloud records. UI shows pending count and sync error reason.'),
        ('Financial integrity', 'Server recalculates/validates totals and permissions. Currency storage uses integer minor units where applicable.'),
        ('Audit data', 'Returns, stock movements, deliveries, price changes, shift close, credit payments, project actions, plan changes and platform actions retain timestamps and actors.'),
        ('Secrets', 'Supabase service-role key and Paystack private keys are Edge Function/server secrets only. Browser uses publishable key.'),
        ('Backups/monitoring', 'Operational process must define database backups, error tracking, alerting, CI/CD checks, incident ownership and recovery testing before scaled rollout.'),
    ], [2600, 6760])
    heading(doc, '12. PM test scenarios')
    table(doc, ['Scenario', 'Pass condition'], [
        ('Offline cash sale', 'Disable network; complete sale; reload; queue remains. Restore network; sale syncs once and appears in cloud sales.'),
        ('Credit sale', 'Enter client, phone, due date and optional deposit; completion clears checkout; credit register reflects balance; partial payment updates history/status.'),
        ('Supplier price increase', 'Receive delivery with new unit cost/date; stock rises; current cost and cost/price history update; later sale profit uses its own snapshot.'),
        ('Partial return', 'Manager selects one sold item and partial quantity; server rejects over-return; inventory restores correct quantity and return actor appears in audit.'),
        ('Historical entry', 'Admin records a past sale with original price; record appears in sales/reports as historical and stock remains unchanged unless correction is selected.'),
        ('Service project', 'Create new client in project wizard; include date/PDF; add payment, expense and stage; invoice prints company branding and payment breakdown.'),
        ('Plan entitlement', 'Starter tenant cannot access storefront. After Growth entitlement, storefront management becomes available. Product limit blocks additional product creation when reached.'),
        ('Mobile platform', 'At mobile width hamburger opens platform sidebar, outside tap/selection closes it, and content remains unobscured.'),
    ], [3000, 6360])
    heading(doc, '13. Operational runbook and unresolved follow-ups')
    bullets(doc, ['Run migrations in timestamp order and deploy Edge Functions after their related database functions/tables exist.', 'For billing, configure Paystack secrets and webhook endpoints in Supabase before enabling upgrades as a production promise.', 'Maintain a tested support path for “Sync needs attention”; never instruct staff to clear IndexedDB before confirming pending transactions are reconciled.', 'Confirm all critical server functions use the actual current role/membership source and write RLS tests for admin, manager and cashier.', 'Before multi-store launch, define branch onboarding, inventory transfer, reporting aggregation and user membership model.'])
    doc.save(OUT / 'Kroniqos_Functional_Specification.docx')


if __name__ == '__main__':
    make_prd()
    make_functional_spec()
    print(OUT / 'Kroniqos_PRD.docx')
    print(OUT / 'Kroniqos_Functional_Specification.docx')
