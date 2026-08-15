from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = 'docs/Naira_POS_PRD.docx'
BLUE = '1F4D78'; LIGHT = 'E8EEF5'; GRAY = 'F2F4F7'; DARK = '0B2545'; RED = '9B1C1C'

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def set_cell(cell, text, bold=False, color=None, size=9.5):
    cell.text = ''
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text)); r.bold = bold; r.font.name = 'Calibri'; r.font.size = Pt(size)
    if color: r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.LEFT; t.style = 'Table Grid'
    for i, h in enumerate(headers):
        set_cell(t.rows[0].cells[i], h, True, DARK); shade(t.rows[0].cells[i], LIGHT)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row): set_cell(cells[i], value)
    if widths:
        for row in t.rows:
            for i, width in enumerate(widths): row.cells[i].width = Inches(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f'Heading {level}'); p.add_run(text)
    return p

def para(doc, text='', bold_lead=None):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6)
    if bold_lead and text.startswith(bold_lead):
        p.add_run(bold_lead).bold = True; p.add_run(text[len(bold_lead):])
    else: p.add_run(text)
    return p

def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet'); p.add_run(item)

def page_break(doc): doc.add_page_break()

doc = Document()
sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(.49)

styles = doc.styles
normal = styles['Normal']; normal.font.name = 'Calibri'; normal.font.size = Pt(11); normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.1
for name, size, color, before, after in [('Heading 1',16,BLUE,16,8),('Heading 2',13,BLUE,12,6),('Heading 3',12,'1F4D78',8,4)]:
    s = styles[name]; s.font.name='Calibri'; s.font.size=Pt(size); s.font.color.rgb=RGBColor.from_string(color); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after)

header = sec.header.paragraphs[0]; header.text = 'NAIRA POS  |  Product Requirements Document'; header.runs[0].font.color.rgb=RGBColor.from_string('666666'); header.runs[0].font.size=Pt(8)
footer = sec.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT; footer.add_run('Confidential — internal project planning')

# Cover
p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(54); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r=p.add_run('PRODUCT REQUIREMENTS DOCUMENT'); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=RGBColor.from_string(BLUE)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(4); r=p.add_run('Naira POS'); r.bold=True; r.font.size=Pt(30); r.font.color.rgb=RGBColor.from_string(DARK)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(22); r=p.add_run('Single-store, offline-resilient supermarket point-of-sale MVP'); r.font.size=Pt(15); r.font.color.rgb=RGBColor.from_string('555555')
table(doc, ['Document control', 'Value'], [
    ('Owner', 'Product Manager'), ('Version', '1.0'), ('Date', '31 July 2026'), ('Target market', 'Nigerian single-store supermarkets'), ('Currency', 'Nigerian naira (₦)'), ('Technology', 'React, TypeScript, Zustand, Dexie, Supabase, Ant Design, Tailwind CSS')], [1.7,4.7])
heading(doc, 'Executive summary', 1)
para(doc, 'Naira POS is a touch-friendly supermarket POS that keeps cashiers selling during network outages while providing managers and administrators with controlled inventory, staff, reporting, and supplier operations. The MVP is designed for one store today, with store-scoped data and role controls that allow branches to be added later.')
para(doc, 'Product decision: checkout is offline-first. Inventory, returns, supplier deliveries, staff, settings, and management reports are online-first because they change shared financial or stock truth.')
page_break(doc)

heading(doc, '1. Product vision and goals')
para(doc, 'Vision: give a Nigerian supermarket a fast, reliable checkout experience without sacrificing stock, cost, and audit accuracy.')
table(doc, ['Goal', 'Success measure'], [
 ('Fast service at till', 'Cashier can scan/add products, hold a sale, choose payment, and complete a local sale during an outage.'),
 ('Protect business data', 'Each sale receives an idempotent queue record and is synced once Supabase confirms it.'),
 ('Controlled operations', 'Only authorised roles can access inventory, pricing, staff, reports, and returns.'),
 ('Management visibility', 'Admin can filter sales/profit by day, month, year, and payment method; export CSV.'),
 ('Accurate unit economics', 'Cost is captured at sale time and supplier deliveries create price audit records.')], [2.1,4.4])
heading(doc, 'Non-goals for MVP', 2)
bullets(doc, ['Multi-store/branch operations, consolidated branch reporting, and inter-branch transfers.', 'Integrated payment processor settlement; card, transfer, and mobile-money are recorded payment methods only.', 'Accounting-system integration, purchase orders, tax filing, loyalty, promotions, and e-commerce.', 'Guaranteed prevention of stock oversell across multiple offline terminals.'])

heading(doc, '2. Personas and role-based access')
table(doc, ['Role', 'Primary jobs', 'Access'], [
 ('Cashier', 'Serve customers quickly; scan items; hold/resume cart; take payment.', 'Checkout, held sales, personal shift operation. No price, stock, supplier, staff, or profit editing.'),
 ('Manager', 'Maintain stock, handle approved returns, receive deliveries, manage daily operations.', 'Inventory changes and supplier deliveries online; sales review and returns. No staff administration or profit report page.'),
 ('Administrator', 'Own the store configuration, staff, pricing, profit, and governance.', 'All manager access plus staff, store settings, profit reports, category creation, product pricing and CSV export.')], [1.1,2.3,3.1])

page_break(doc)
heading(doc, '3. Operating model: offline-first vs online-first')
table(doc, ['Workflow', 'Mode', 'Reason'], [
 ('Checkout / cart / held sales / receipt', 'Offline-first', 'Customer service must continue through weak or lost internet.'),
 ('Product catalogue read', 'Local cache + online refresh', 'Checkout is fast; current product data is fetched when connected.'),
 ('Sales sync', 'Background outbox', 'Immutable sale event is replayed with a client UUID to avoid duplicates.'),
 ('Stock counts, stock adjustments, deliveries, cost/price changes', 'Online-first', 'These are shared stock and financial changes; server is authoritative.'),
 ('Returns/exchanges', 'Online-first target', 'Return changes stock and financial audit; requires manager authorisation.'),
 ('Shifts, staff, settings, reports', 'Online-first', 'These require current shared truth and a reliable audit trail.')], [2.1,1.3,3.1])
para(doc, 'Operational caveat: if multiple terminals are offline, each can sell the last unit. The server validates stock on sync; the later conflicting sale requires manager reconciliation.')
page_break(doc)

heading(doc, '4. Functional requirements')
heading(doc, '4.1 Checkout and payment', 2)
bullets(doc, ['Search by product name or SKU; USB/Bluetooth barcode scanner behaves as keyboard input and adds an exact matching product when it sends Enter.', 'Touch-optimised product cards and large payment/quantity controls.', 'Cash, card/POS, bank transfer, and mobile money payment methods.', 'Cash tendered and change due calculation; insufficient tender blocks completion.', 'Manager/admin discount control. Server validates final total and records approved discounts.', 'Hold sale clears the active cart and stores it locally; Resume restores the cart and payment choice; manager/cashier can discard held cart.', 'Receipt modal supports receipt-only printing and later receipt viewing from Sales.'])
heading(doc, '4.2 Product and inventory', 2)
bullets(doc, ['Add product with name, scanned/typed barcode, category, cost price, selling price, and opening stock.', 'Choose existing category or create a new category during product creation.', 'Edit product name, barcode, cost, selling price, availability, and category data as authorised.', 'Low-stock indicator follows the configurable threshold.', 'Stock count records physical quantity and calculates the variance.', 'Stock adjustment records delivery, correction, damaged/waste, or count reasons.', 'Supplier delivery records supplier, product, quantity, and unit cost; it increases stock and updates current cost price.', 'Inventory movement history records sale, return, delivery, correction, damage, and stock count events.'])
heading(doc, '4.3 Sales, returns and reports', 2)
bullets(doc, ['Sales page includes filters, recent sales, receipt viewing, return-item selection, return activity, and CSV export for authorised users.', 'Partial return lets managers select individual items and quantities; stock restoration and return audit are required.', 'Profit Reports is an admin page with day/month/year and payment filters, net revenue, COGS, gross profit, margin, per-receipt quantity, cost, and profit.', 'Current cost is captured on each sale item so later cost changes do not rewrite historic gross profit.'])
heading(doc, '4.4 Staff and settings', 2)
bullets(doc, ['Supabase Auth handles sign-in/sign-out. Profiles carry store membership and app role.', 'Admin-only Staff page creates/manages staff via the secure manage-staff Edge Function.', 'Store settings cover name, address, phone, receipt footer, VAT, low-stock threshold, and enabled payment methods.'])

heading(doc, '5. Core user flows', 1)
table(doc, ['Flow', 'Happy path', 'Exception / rule'], [
 ('Cash sale', 'Scan/add → adjust quantity → choose Cash → enter tendered amount → complete → local receipt → sync.', 'Offline sale remains pending until server confirmation; must not be deleted by reset.'),
 ('Unknown barcode', 'Manager/admin scans unknown code → Add Product opens with barcode prefilled → enter category/cost/sell/stock → save online.', 'Cashier sees not-found message; cannot create catalogue records.'),
 ('Receive delivery', 'Inventory → Receive delivery → choose product, supplier, qty, unit cost → server confirms → local cache refreshes.', 'Requires online manager/admin session. No offline delivery queue.'),
 ('Stock count', 'Inventory → Count → enter physical quantity → server records delta and movement.', 'Requires online manager/admin session.'),
 ('Return item', 'Sales → Return items → select item quantities → manager confirms → server audits and restores stock.', 'Must prevent quantity greater than originally sold and require online approval.'),
 ('Report', 'Admin → Profit reports → choose period/payment → review totals and per-receipt COGS/profit → export CSV.', 'Remote sales without sale-item cost detail should be marked as incomplete until full item pull exists.')], [1.15,3.15,2.2])
page_break(doc)

heading(doc, '6. Data and integration requirements')
table(doc, ['Entity', 'Key fields', 'Notes'], [
 ('profiles', 'id, store_id, full_name, role', 'Linked to Supabase Auth; role is authoritative server-side.'),
 ('stores', 'id, name, currency_code', 'Store scope is present now to support branches later.'),
 ('products', 'store_id, category_id, sku, price_kobo, cost_price_kobo, stock_quantity, active', 'All money is stored in kobo; UI displays ₦.'),
 ('sales / sale_items', 'receipt, cashier, payment, totals, unit_price, cost snapshot', 'Sale IDs provide sync idempotency.'),
 ('stock_movements', 'product, quantity delta, reason, operation ID', 'Immutable inventory ledger.'),
 ('sale_returns / return items', 'sale, quantities, returned_by, operation ID', 'Manager audit and stock restoration.'),
 ('suppliers / supplier_deliveries', 'supplier, product, qty, unit cost, received_by', 'Delivery and supplier cost audit.'),
 ('product_price_history', 'old cost, new cost, changed_by, changed_at', 'Tracks supplier-driven cost increases.'),
 ('cash_shifts', 'opening, expected, counted, variance', 'Online confirmed cash-control record.')], [1.55,2.8,2.15])
heading(doc, 'API and security requirements', 2)
bullets(doc, ['Supabase Row Level Security limits every store-scoped table to the current staff member’s store.', 'Server-side RPCs validate role and store for record_sale, adjust_stock, returns, supplier deliveries, and cash shifts.', 'The browser never receives the Supabase service-role key.', 'The manage-staff Edge Function is admin-only and is the only supported browser path for staff creation.', 'All server mutations require authenticated session, idempotency key where retries are possible, and auditable actor/timestamp fields.'])

page_break(doc)
heading(doc, '7. Non-functional requirements')
table(doc, ['Area', 'Requirement'], [
 ('Performance', 'Checkout interaction should feel immediate from local Dexie cache; product refresh must not block cart operations.'),
 ('Reliability', 'Pending checkout queue survives refresh/restart on the same device; destructive reset is blocked while pending records exist.'),
 ('Usability', 'Touch controls at least 52px; primary checkout control at least 64px; clear status messages.'),
 ('Security', 'Least-privilege RBAC and server enforcement; no role switching in production UI.'),
 ('Auditability', 'Stock, return, price, supplier, shift, and staff actions include time and actor wherever server data is available.'),
 ('Observability', 'Sync status shows pending count and actionable error reason; staff can retry but cannot silently discard business records.'),
 ('Currency', 'Use ₦ on screen and integer kobo in database logic.')], [1.5,5.0])

heading(doc, '8. Acceptance criteria')
bullets(doc, ['Cashier can complete a sale offline, close/reopen browser, and see it pending; upon reconnect it syncs exactly once.', 'Admin/manager can receive supplier delivery online and see stock and cost update; price history records cost change.', 'Cashier cannot change stock, cost, category, staff, or profit data even if they manipulate client UI.', 'Admin can filter profit report by day, month, year, and payment method and export the same filtered sales to CSV.', 'Return cannot exceed sold quantity; manager identity is recorded server-side.', 'Inventory mutation while offline is blocked without changing local shared stock.', 'Sales/refunds/stock sync failures name the failing operation and preserve the unsynced checkout record.'])

heading(doc, '9. Delivery plan and priorities')
table(doc, ['Phase', 'Scope', 'Exit criteria'], [
 ('MVP checkout', 'Auth/RBAC, catalogue cache, barcode checkout, payments, held sales, receipts, local sale queue.', 'Cashier can sell reliably with/without network.'),
 ('Management', 'Inventory, stock ledger, cash shifts, staff, reports, CSV, returns.', 'Managers operate daily store controls with audit records.'),
 ('Unit economics', 'Cost/selling price, supplier deliveries, cost snapshots, profit report, price history.', 'Admin can explain gross profit by receipt and period.'),
 ('Hardening', 'Online-first returns/settings, server-side report item pull, role test suite, backup/monitoring, reconciliation workflow.', 'No unresolved shared-state conflicts or invisible sync failures.'),
 ('Future', 'Branches, transfers, purchase orders, accounting export, promotions, loyalty, payment integrations.', 'Validated after single-store adoption.')], [1.15,3.15,2.2])

heading(doc, '10. Risks and decisions for the project manager')
table(doc, ['Risk / decision', 'Impact', 'Mitigation / owner'], [
 ('Inventory permission mismatch', 'Managers cannot operate stock controls.', 'Verify Supabase profile/store mapping and RPC grants before rollout. Engineering + PM.'),
 ('Offline oversell', 'Later sale may fail server stock validation.', 'Manager reconciliation policy; avoid multiple offline terminals for scarce SKUs. Operations.'),
 ('Stale pricing at checkout', 'Server may reject sale total after a price change.', 'Refresh catalogue on connection; price-change communication; manager handling for disputed queued sale. Product/Ops.'),
 ('Historical remote COGS gaps', 'Cross-device profit report may be incomplete.', 'Pull sale-item cost details from Supabase before multi-terminal reporting launch. Engineering.'),
 ('Migration order', 'Features fail if required RPC/table is missing.', 'Maintain ordered migration checklist and run in staging first. PM/Engineering.'),
 ('Data reset', 'Unsynced checkout data could be lost.', 'No destructive reset with pending queue; backups and user training. Operations.')], [1.7,2.4,2.4])

heading(doc, '11. Release readiness checklist')
bullets(doc, ['All Supabase migrations applied in order and Edge Function deployed.', 'At least one admin, one manager, and one cashier tested with their real store profiles.', 'Online-first inventory and return actions verified against Supabase with an admin account.', 'Offline checkout test completed: sale, held sale, reconnect, sync, and receipt verification.', 'Supplier delivery, cost update, and profit report validated with test product.', 'CSV exported and opened in Excel; ₦ values and columns checked.', 'Backups, monitoring, support contact, and staff training plan agreed.'])

doc.save(OUT)
print(OUT)
